# -*- coding: utf-8 -*-
from django.db.models import Q
from docx import Document
from publications.models import Publication, Journal


def export_publication_to_doc(queryset, document=None):
    if not document:
        document = Document()
    document.add_heading('Экспорт публикаций', 0)
    document.add_heading('Статьи в изданиях, рекомендованных ВАК и/или входящих'
                         ' в международные базы цитирования WoS и Scopus:', level=2)
    articles = queryset.filter(type=Publication.ARTICLE).\
        filter(Q(is_rinc=True) | Q(is_vak=True) | Q(is_wos=True) | Q(is_scopus=True))
    for article in articles:
        document.add_paragraph(
            article.get_harvard(), style='List Number'
        )

    document.add_heading('Статьи в трудах конференций:', level=2)
    proceedings = queryset.filter(type=Publication.PROCEEDINGS)
    for p in proceedings:
        document.add_paragraph(
            p.get_harvard(), style='List Number'
        )

    document.add_heading('Авторские монографии:', level=2)
    mono = queryset.filter(type=Publication.MONOGRAPH)
    for m in mono:
        document.add_paragraph(
            m.get_harvard(), style='List Number'
        )

    document.add_heading('Коллективные монографии (глава):', level=2)
    g_mono = queryset.filter(type=Publication.GROUP_MONOGRAPH)
    for gm in g_mono:
        document.add_paragraph(
            gm.get_harvard(), style='List Number'
        )

    document.add_heading('Тезисы конференций:', level=2)
    confs = queryset.filter(type=Publication.THESES_CONFERENCE)
    for c in confs:
        document.add_paragraph(
            c.get_harvard(), style='List Number'
        )

    document.add_heading('Учебно-методические материалы:', level=2)
    teach_materials = queryset.filter(type=Publication.TEACHING_MATERIALS)
    for t in teach_materials:
        document.add_paragraph(
            t.get_harvard(), style='List Number'
        )

    document.add_heading('Прочие статьи:', level=2)
    another_articles = queryset.filter(type=Publication.ARTICLE). \
        filter(Q(is_other_db=True) & Q(Q(is_rinc=False) | Q(is_vak=False) | Q(is_wos=False) | Q(is_scopus=False)))
    for article in another_articles:
        document.add_paragraph(
            article.get_harvard(), style='List Number'
        )

    document.add_heading('Авторские свидетельства:', level=2)
    patents = queryset.filter(Q(type=Publication.PATENT) | Q(type=Publication.PATENT_BD))
    for patent in patents:
        document.add_paragraph(
            patent.get_harvard(), style='List Number'
        )

    return document


def export_grants_to_doc(queryset, document=None):
    if not document:
        document = Document()
    document.add_heading('Экспорт грантов', 0)

    for grant in queryset:
        document.add_paragraph(
            grant.export(), style='List Number'
        )

    return document


def export_conference_to_doc(queryset, document=None):
    if not document:
        document = Document()

    document.add_heading('Экспорт конференций', 0)

    national_conferences = queryset.filter(publication__journal__conf_type=Journal.NATIONAL)
    national_count = national_conferences.count()
    if national_conferences:
        document.add_heading('Отечественные мероприятия:', level=2)
        document.add_paragraph('- отечественные мероприятия: {}'.format(national_count))

        table = document.add_table(rows=1, cols=3)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название мероприятия'
        hdr_cells[1].text = 'Место и время проведения'
        hdr_cells[2].text = 'Название доклада'
        for conf in national_conferences:
            row_cells = table.add_row().cells
            row_cells[0].text = conf.publication.journal.name_ru
            row_cells[1].text = conf.publication.journal.get_place_dates()
            row_cells[2].text = conf.publication.title

    international_conferences = queryset.filter(publication__journal__conf_type=Journal.INTERNATIONAL)
    inter_count = international_conferences.count()
    if international_conferences:
        document.add_heading('Зарубежные мероприятия:', level=2)
        document.add_paragraph('- зарубежные мероприятия: {}'.format(inter_count))

        table = document.add_table(rows=1, cols=3)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Название мероприятия'
        hdr_cells[1].text = 'Место и время проведения'
        hdr_cells[2].text = 'Название доклада'
        for conf in international_conferences:
            row_cells = table.add_row().cells
            row_cells[0].text = conf.publication.journal.name_ru
            row_cells[1].text = conf.publication.journal.get_place_dates()
            row_cells[2].text = conf.publication.title

    return document


def export_from_profile(publications, grants_member, grants_head, conference):
    document = Document()

    if publications:
        document = export_publication_to_doc(publications, document)
    if grants_member:
        document = export_grants_to_doc(grants_member, document)
    if grants_head:
        document = export_grants_to_doc(grants_head, document)
    if conference:
        document = export_conference_to_doc(conference, document)

    return document
